import os
import sys
try:
    from docx import Document
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx>=1.2.0")
    sys.exit(1)

from docx_annotator import add_comment_to_doc

def test_annotator():
    # Setup: Create a dummy doc
    test_file = "test_doc.docx"
    doc = Document()
    doc.add_paragraph("This is a legal contract with some risk terms.")
    doc.add_paragraph("Another paragraph.")
    doc.save(test_file)
    print(f"Created {test_file}")

    try:
        # Run Annotator
        target = "risk terms"
        comment = "Please define this term clearly."
        output = add_comment_to_doc(test_file, target, comment)

        # Assertions
        if output and os.path.exists(output):
            print("SUCCESS: Output file created.")
            # Verify content logic if needed (requires reloading doc)
            # For now, file existence and script non-failure is the basic test.
        else:
            print("FAILURE: Output file not created.")

    except Exception as e:
        print(f"FAILURE: Exception occurred: {e}")

    finally:
        # Cleanup
        if os.path.exists(test_file):
            os.remove(test_file)
        # Optional: remove output file
        # if output and os.path.exists(output): os.remove(output)

if __name__ == "__main__":
    test_annotator()
